#!/usr/bin/env python3
"""Generate the undergraduate thesis deliverables from the current draft.

The script produces:
- final_thesis.html
- final_thesis.docx, if LibreOffice is available
- final_thesis.pdf, if LibreOffice can export it
"""

from __future__ import annotations

import html
import re
import subprocess
from pathlib import Path

try:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Inches, Pt, RGBColor
except Exception:  # pragma: no cover - optional generation path
    Document = None
    WD_SECTION = None
    WD_CELL_VERTICAL_ALIGNMENT = None
    WD_TABLE_ALIGNMENT = None
    WD_ALIGN_PARAGRAPH = None
    WD_TAB_ALIGNMENT = None
    WD_TAB_LEADER = None
    OxmlElement = None
    qn = None
    Cm = None
    Inches = None
    Pt = None
    RGBColor = None


ROOT = Path(__file__).resolve().parents[2]
THESIS_DIR = ROOT / "docs" / "thesis"
DRAFT = THESIS_DIR / "thesis_draft.md"
OUT_HTML = THESIS_DIR / "final_thesis.html"
OUT_ODT = THESIS_DIR / "final_thesis.odt"
OUT_DOCX = THESIS_DIR / "final_thesis.docx"
OUT_PDF = THESIS_DIR / "final_thesis.pdf"

EAST_ASIA_FONT = "SimSun"
HEADING_FONT = "SimHei"
LATIN_FONT = "Times New Roman"
MATH_FONT = "Cambria Math"
CITATION_RE = re.compile(r"\[\d+(?:-\d+)?\]")
CAPTION_RE = re.compile(r"^(图|表)\s*\d+\s*(.*)$")


COVER = {
    "title_cn": "基于 KAN 超图的药物协同预测研究",
    "title_en": "KAN-Enhanced Hypergraph Neural Networks for Drug Synergy Prediction",
    "name": "赵新宇",
    "student_id": "202200130116",
    "college": "计算机科学与技术学院",
    "major": "计算机科学与技术",
    "grade": "2022级",
    "advisor": "崔学峰",
    "date": "2026 年 5 月 10 日",
}


EN_ABSTRACT = """
Drug combination therapy is an important strategy for cancer treatment, but exhaustive wet-lab screening of all possible drug pairs, dose settings, and cellular contexts is prohibitively expensive. This thesis studies how Kolmogorov-Arnold Networks (KANs) can be used in drug synergy prediction. Initial experiments show that directly replacing the final multilayer perceptron with a KAN prediction head does not consistently improve performance, suggesting that KAN is not a universal drop-in replacement for high-dimensional fused representations. Based on this observation, this work investigates a more structured design: combining KAN with molecular graph encoders, hypergraph message passing, and drug-pair aggregation graphs.

Using DrugComb-derived drug combination data, the study evaluates classification and Loewe-score regression tasks under Random split and Cold-drug split settings. The results indicate that Random split mainly reflects interpolation over known drugs, while Cold-drug split is more relevant to realistic screening of unseen compounds. Among the explored variants, HgKAN-Agg, which first aggregates two drug embeddings into a drug-pair node and then performs KAN-enhanced graph message passing with the cell node, is the only main KAN variant that improves both classification and regression metrics over the MLP baseline in the Cold-drug split setting. In contrast, direct KAN heads, KAN hypergraph variants, and stacking DrugKAN with HgKAN-Agg do not always improve performance, showing that the placement of KAN is crucial.

Beyond predictive accuracy, this thesis builds an interpretability chain consisting of KAN edge-function visualization, drug perturbation analysis, drug-pair retention response curves, and atom-level saliency. A Cold-drug split case study on Lenalidomide and mitomycin C in the IGROV1 cell line shows that the model prediction depends strongly on drug-pair information, and the highlighted molecular regions overlap with known pharmacologically relevant substructures. Overall, the results support a cautious conclusion: in drug synergy prediction, KAN is more suitable as a learnable nonlinear message function inside structured graph or hypergraph modules than as an unstructured replacement for an MLP head.
""".strip()


EN_KEYWORDS = (
    "drug synergy prediction, graph neural network, hypergraph neural network, "
    "Kolmogorov-Arnold Network, interpretable machine learning"
)


# Kept only as a fallback template; generated outputs use the numbered
# references in thesis_draft.md via extract_references().
REFERENCES = [
    "Liu Z, Wang Y, Vaidya S, Ruehle F, Halverson J, Soljacic M, Hou T Y, Tegmark M. KAN: Kolmogorov-Arnold Networks. arXiv:2404.19756, 2024.",
    "Li L, Zhang Y, Wang G, Xia K. Kolmogorov-Arnold graph neural networks for molecular property prediction. Nature Machine Intelligence, 2025, 7:1346-1354. DOI: 10.1038/s42256-025-01087-7.",
    "Preuer K, Lewis R P I, Hochreiter S, Bender A, Bulusu K C, Klambauer G. DeepSynergy: predicting anti-cancer drug synergy with Deep Learning. Bioinformatics, 2018, 34(9):1538-1546. DOI: 10.1093/bioinformatics/btx806.",
    "Kuru H I, Tastan O, Cicek A E. MatchMaker: A Deep Learning Framework for Drug Synergy Prediction. IEEE/ACM Transactions on Computational Biology and Bioinformatics, 2022, 19(4):2334-2344. DOI: 10.1109/TCBB.2021.3086702.",
    "Schwarz K, Pliego-Mendieta A, Mollaysa A, Planas-Paz L, Pauli C, Allam A, Krauthammer M. DDoS: A Graph Neural Network based Drug Synergy Prediction Algorithm. Proceedings of Machine Learning Research, CHIL, 2024; arXiv:2210.00802.",
    "Zagidullin B, Aldahdooh J, Zheng S, Wang W, Wang Y, Saad J, Malyutina A, Jafari M, Tanoli Z, Pessia A, Tang J. DrugComb: an integrative cancer drug combination data portal. Nucleic Acids Research, 2019, 47(W1):W43-W51. DOI: 10.1093/nar/gkz337.",
    "Zheng S, Aldahdooh J, Shadbahr T, Wang Y, Aldahdooh D, Bao J, Wang W, Tang J. DrugComb update: a more comprehensive drug sensitivity data repository and analysis portal. Nucleic Acids Research, 2021, 49(W1):W174-W184. DOI: 10.1093/nar/gkab438.",
    "Holbeck S L, Camalier R, Crowell J A, Govindharajulu J P, Hollingshead M, Anderson L W, Polley E, Rubinstein L. The National Cancer Institute ALMANAC: A Comprehensive Screening Resource for the Detection of Anticancer Drug Pairs with Enhanced Therapeutic Activity. Cancer Research, 2017, 77(13):3564-3576. DOI: 10.1158/0008-5472.CAN-17-0489.",
    "Ianevski A, Giri A K, Aittokallio T. SynergyFinder 2.0: visual analytics of multi-drug combination synergies. Nucleic Acids Research, 2020, 48(W1):W488-W493. DOI: 10.1093/nar/gkaa216.",
    "Ianevski A, He L, Aittokallio T, Tang J. SynergyFinder: a web application for analyzing drug combination dose-response matrix data. Bioinformatics, 2020, 36(8):2645. DOI: 10.1093/bioinformatics/btaa102.",
    "Wooten D J, Meyer C T, Lubbock A L R, Quaranta V, Lopez C F. MuSyC is a consensus framework that unifies multi-drug synergy metrics for combinatorial drug discovery. Nature Communications, 2021, 12:4607. DOI: 10.1038/s41467-021-24789-z.",
    "Liu X, Song C, Liu S, Li M, Zhou X, Zhang W. Multi-way relation-enhanced hypergraph representation learning for anti-cancer drug synergy prediction. Bioinformatics, 2022, 38(20):4782-4789. DOI: 10.1093/bioinformatics/btac579.",
    "Menden M P, Iorio F, Garnett M, McDermott U, Benes C H, Ballester P J, Saez-Rodriguez J. Machine Learning Prediction of Cancer Cell Sensitivity to Drugs Based on Genomic and Chemical Properties. PLoS ONE, 2013, 8(4):e61318. DOI: 10.1371/journal.pone.0061318.",
    "Garnett M J, Edelman E J, Heidorn S J, Greenman C D, Dastur A, Lau K W, Greninger P, Thompson I R, et al. Systematic identification of genomic markers of drug sensitivity in cancer cells. Nature, 2012, 483(7391):570-575. DOI: 10.1038/nature11005.",
    "Iorio F, Knijnenburg T A, Vis D J, Bignell G R, Menden M P, Schubert M, Aben N, Goncalves E, et al. A Landscape of Pharmacogenomic Interactions in Cancer. Cell, 2016, 166(3):740-754. DOI: 10.1016/j.cell.2016.06.017.",
    "Kipf T N, Welling M. Semi-Supervised Classification with Graph Convolutional Networks. arXiv:1609.02907, 2016.",
    "Gilmer J, Schoenholz S S, Riley P F, Vinyals O, Dahl G E. Neural Message Passing for Quantum Chemistry. arXiv:1704.01212, 2017.",
    "Velickovic P, Cucurull G, Casanova A, Romero A, Lio P, Bengio Y. Graph Attention Networks. arXiv:1710.10903, 2017.",
    "Xu K, Hu W, Leskovec J, Jegelka S. How Powerful are Graph Neural Networks? arXiv:1810.00826, 2018.",
    "Fey M, Lenssen J E. Fast Graph Representation Learning with PyTorch Geometric. arXiv:1903.02428, 2019.",
    "Hu W, Fey M, Zitnik M, Dong Y, Ren H, Liu B, Catasta M, Leskovec J. Open Graph Benchmark: Datasets for Machine Learning on Graphs. arXiv:2005.00687, 2020.",
    "Chamberlain P P, Lopez-Girona A, Miller K, Carmel G, Pagarigan B, Chie-Leon B, Rychak E, Corral L G, et al. Structure of the human Cereblon-DDB1-lenalidomide complex reveals basis for responsiveness to thalidomide analogs. Nature Structural & Molecular Biology, 2014, 21(9):803-809. DOI: 10.1038/nsmb.2874.",
    "Tomasz M. Mitomycin C: small, fast and deadly (but very selective). Chemistry & Biology, 1995, 2(9):575-579. DOI: 10.1016/1074-5521(95)90120-5.",
    "Bliss C I. The toxicity of poisons applied jointly. Annals of Applied Biology, 1939, 26(3):585-615. DOI: 10.1111/j.1744-7348.1939.tb06990.x.",
]


def split_draft(text: str) -> tuple[str, str, str, str]:
    abstract_match = re.search(r"## 摘要\s+(.*?)\n\n关键词：(.+?)\n\n## 1 ", text, re.S)
    if not abstract_match:
        raise RuntimeError("Cannot locate Chinese abstract in thesis_draft.md")
    abstract = abstract_match.group(1).strip()
    keywords = abstract_match.group(2).strip()

    body_start = text.index("## 1 ")
    refs_start = text.index("## 参考文献")
    body = text[body_start:refs_start].strip()
    return COVER["title_cn"], abstract, keywords, body


def extract_references(text: str) -> list[str]:
    refs_text = text.split("## 参考文献", 1)[1]
    refs: list[str] = []
    current: list[str] = []
    for raw_line in refs_text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        match = re.match(r"^\[\d+\]\s*(.+)", line)
        if match:
            if current:
                refs.append(" ".join(current).strip())
            current = [match.group(1).strip()]
        elif current:
            current.append(line)
    if current:
        refs.append(" ".join(current).strip())
    return refs


def esc(text: str) -> str:
    return html.escape(text, quote=False)


def citation_href(citation_text: str) -> str:
    match = re.match(r"\[(\d+)", citation_text)
    return f"ref-{match.group(1)}" if match else "refs"


def esc_inline(text: str) -> str:
    out: list[str] = []
    pos = 0
    for match in INLINE_MATH_RE.finditer(text):
        if match.start() > pos:
            out.append(esc_text_with_citations(text[pos:match.start()]))
        out.append(render_math_html(match.group(1)))
        pos = match.end()
    if pos < len(text):
        out.append(esc_text_with_citations(text[pos:]))
    return "".join(out)


def esc_text_with_citations(text: str) -> str:
    escaped = esc(text)
    return CITATION_RE.sub(
        lambda m: f"<a class='citation' href='#{citation_href(m.group(0))}'><sup>{m.group(0)}</sup></a>",
        escaped,
    )


def render_table(lines: list[str]) -> str:
    rows = []
    for line in lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", c or "") for c in cells):
            continue
        rows.append(cells)
    if not rows:
        return ""
    head, *body = rows
    out = ["<table>", "<thead><tr>"]
    out += [f"<th>{esc(c)}</th>" for c in head]
    out += ["</tr></thead><tbody>"]
    for row in body:
        out.append("<tr>")
        out += [f"<td>{esc(c)}</td>" for c in row]
        out.append("</tr>")
    out.append("</tbody></table>")
    return "\n".join(out)


def render_markdown(md: str) -> str:
    out: list[str] = []
    lines = md.splitlines()
    i = 0
    paragraph: list[str] = []
    pending_figure = False
    expect_figure_note = False

    def flush_paragraph() -> None:
        if paragraph:
            text = " ".join(x.strip() for x in paragraph).strip()
            out.append(f"<p>{esc_inline(text)}</p>")
            paragraph.clear()

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            flush_paragraph()
            i += 1
            continue

        if is_caption_line(stripped, "表") and next_nonempty_line(lines, i + 1).startswith("|"):
            flush_paragraph()
            out.append(f"<p class='caption'>{esc_inline(stripped)}</p>")
            i += 1
            continue

        if is_caption_line(stripped, "图") and pending_figure:
            flush_paragraph()
            out.append(f"<p class='caption'>{esc_inline(stripped)}</p>")
            pending_figure = False
            expect_figure_note = True
            i += 1
            continue

        if expect_figure_note:
            if is_markdown_block_boundary(stripped, lines, i):
                expect_figure_note = False
            else:
                flush_paragraph()
                note_lines = []
                while i < len(lines):
                    candidate = lines[i].strip()
                    if not candidate or is_markdown_block_boundary(candidate, lines, i):
                        break
                    note_lines.append(lines[i])
                    i += 1
                note = ensure_note_prefix(" ".join(x.strip() for x in note_lines).strip())
                out.append(f"<p class='figure-note'>{esc_inline(note)}</p>")
                expect_figure_note = False
                continue

        if stripped.startswith("|"):
            flush_paragraph()
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            out.append(render_table(table_lines))
            continue

        image_match = re.match(r"!\[(.*?)\]\((.*?)\)", stripped)
        if image_match:
            flush_paragraph()
            alt, rel = image_match.groups()
            src = (THESIS_DIR / rel).resolve().as_uri()
            out.append(
                f'<figure><img src="{src}" alt="{esc(alt)}"/></figure>'
            )
            pending_figure = True
            i += 1
            continue

        if stripped.startswith("### "):
            flush_paragraph()
            out.append(f"<h2>{esc(stripped[4:])}</h2>")
            i += 1
            continue

        if stripped.startswith("## "):
            flush_paragraph()
            out.append(f"<h1>{esc(stripped[3:])}</h1>")
            i += 1
            continue

        if stripped.startswith("- "):
            flush_paragraph()
            items = []
            while i < len(lines) and lines[i].strip().startswith("- "):
                items.append(lines[i].strip()[2:])
                i += 1
            out.append("<ul>" + "".join(f"<li>{esc_inline(x)}</li>" for x in items) + "</ul>")
            continue

        if re.match(r"\d+\.\s+", stripped):
            flush_paragraph()
            items = []
            while i < len(lines) and re.match(r"\d+\.\s+", lines[i].strip()):
                items.append(re.sub(r"^\d+\.\s+", "", lines[i].strip()))
                i += 1
            out.append("<ol>" + "".join(f"<li>{esc_inline(x)}</li>" for x in items) + "</ol>")
            continue

        if is_math_block_start(stripped):
            flush_paragraph()
            equation_lines = []
            i += 1
            while i < len(lines):
                if is_math_block_end(lines[i].strip()):
                    i += 1
                    break
                equation_lines.append(lines[i].strip())
                i += 1
            eq = " ".join(x for x in equation_lines if x)
            out.append(f"<div class='equation'>{render_math_html(eq)}</div>")
            continue

        paragraph.append(line)
        i += 1

    flush_paragraph()
    return "\n".join(out)


def build_toc(body: str) -> str:
    items = []
    for line in body.splitlines():
        if line.startswith("## "):
            title = line[3:].strip()
            items.append(f"<li class='toc-l1'>{esc(title)}</li>")
        elif line.startswith("### "):
            title = line[4:].strip()
            items.append(f"<li class='toc-l2'>{esc(title)}</li>")
    items.append("<li class='toc-l1'>参考文献</li>")
    items.append("<li class='toc-l1'>致谢</li>")
    return "<ul class='toc'>" + "\n".join(items) + "</ul>"


def build_html() -> str:
    draft_text = DRAFT.read_text(encoding="utf-8")
    _, abstract_cn, keywords_cn, body_md = split_draft(draft_text)
    references = extract_references(draft_text)
    body_html = render_markdown(body_md)
    toc_html = build_toc(body_md)
    refs_html = "\n".join(
        f'<li id="ref-{i}">{esc(ref)}</li>' for i, ref in enumerate(references, 1)
    )

    return f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8"/>
  <title>{esc(COVER['title_cn'])}</title>
  <style>
    @page {{
      size: A4;
      margin: 2.5cm 3.0cm;
    }}
    body {{
      font-family: SimSun, "Noto Serif CJK SC", "Songti SC", serif;
      font-size: 12pt;
      line-height: 1.5;
      color: #111;
    }}
    .page {{
      page-break-after: always;
    }}
    .cover {{
      text-align: center;
      padding-top: 1.5cm;
    }}
    .cover-main {{
      font-family: SimHei, "Noto Sans CJK SC", sans-serif;
      font-size: 26pt;
      font-weight: bold;
      letter-spacing: 0.25em;
      margin-top: 1.2cm;
      margin-bottom: 2.0cm;
    }}
    .cover-title-label {{
      font-size: 14pt;
      margin-bottom: 0.3cm;
    }}
    .cover-title {{
      font-family: SimHei, "Noto Sans CJK SC", sans-serif;
      font-size: 18pt;
      font-weight: bold;
      line-height: 1.4;
      margin-bottom: 0.3cm;
    }}
    .cover-title-en {{
      font-family: "Times New Roman", serif;
      font-size: 14pt;
      font-weight: bold;
      margin-bottom: 2.1cm;
    }}
    .cover-table {{
      margin-left: auto;
      margin-right: auto;
      border-collapse: collapse;
      font-size: 14pt;
      line-height: 2.0;
    }}
    .cover-table td {{
      border: none;
      padding: 0.03cm 0.15cm;
      text-align: left;
    }}
    .cover-table .label {{
      width: 2.8cm;
      text-align: right;
    }}
    .cover-table .value {{
      min-width: 7.8cm;
      border-bottom: 1px solid #111;
      text-align: center;
    }}
    .cover-date {{
      margin-top: 2.0cm;
      font-size: 14pt;
    }}
    .abstract-title, .toc-title, .ref-title, .ack-title, .grade-title {{
      font-family: SimHei, "Noto Sans CJK SC", sans-serif;
      text-align: center;
      font-size: 18pt;
      font-weight: bold;
      margin-top: 0.5cm;
      margin-bottom: 0.8cm;
    }}
    .grade-table {{
      width: 100%;
      border-collapse: collapse;
      font-size: 12pt;
      line-height: 1.5;
    }}
    .grade-table td {{
      border: 1px solid #111;
      padding: 0.3cm;
      height: 2.2cm;
      vertical-align: top;
    }}
    .grade-table .short {{
      height: 1.1cm;
    }}
    .abstract p, .ack p {{
      text-indent: 2em;
      margin: 0 0 0.35cm 0;
    }}
    .keywords {{
      margin-top: 0.45cm;
      text-indent: 0;
    }}
    .keywords b {{
      font-family: SimHei, "Noto Sans CJK SC", sans-serif;
    }}
    h1 {{
      font-family: SimHei, "Noto Sans CJK SC", sans-serif;
      font-size: 16pt;
      font-weight: bold;
      margin-top: 0.6cm;
      margin-bottom: 0.3cm;
    }}
    h2 {{
      font-family: SimHei, "Noto Sans CJK SC", sans-serif;
      font-size: 14pt;
      font-weight: bold;
      margin-top: 0.45cm;
      margin-bottom: 0.25cm;
    }}
    p {{
      text-indent: 2em;
      margin: 0 0 0.25cm 0;
    }}
    ul, ol {{
      margin-top: 0.15cm;
      margin-bottom: 0.25cm;
    }}
    li {{
      margin-bottom: 0.1cm;
    }}
    .toc {{
      list-style: none;
      padding-left: 0;
      font-size: 12pt;
      line-height: 1.8;
    }}
    .toc-l1 {{
      font-weight: bold;
    }}
    .toc-l2 {{
      padding-left: 1.0cm;
      font-weight: normal;
    }}
    table {{
      width: 100%;
      border-collapse: collapse;
      margin: 0.35cm 0;
      font-size: 10.5pt;
      line-height: 1.25;
    }}
    th {{
      border-top: 1.4px solid #111;
      border-bottom: 0.8px solid #111;
      padding: 0.12cm;
      text-align: center;
      font-weight: bold;
    }}
    td {{
      border-bottom: 0.5px solid #888;
      padding: 0.10cm;
      vertical-align: top;
    }}
    .caption {{
      text-align: center;
      text-indent: 0;
      font-size: 10.5pt;
      font-weight: bold;
      margin-top: 0.15cm;
      margin-bottom: 0.15cm;
    }}
    .figure-note {{
      text-indent: 0;
      font-size: 10.5pt;
      line-height: 1.15;
      margin: -0.02cm 0.74cm 0.25cm 0.74cm;
    }}
    .citation {{
      color: inherit;
      text-decoration: none;
    }}
    figure {{
      margin: 0.4cm auto;
      text-align: center;
      page-break-inside: avoid;
    }}
    figure img {{
      max-width: 100%;
      max-height: 17cm;
      object-fit: contain;
    }}
    .equation {{
      text-align: center;
      font-family: "Times New Roman", serif;
      white-space: pre-wrap;
      margin: 0.25cm 0;
    }}
    .references {{
      font-size: 10.5pt;
      line-height: 1.15;
    }}
    .references li {{
      margin-bottom: 0.18cm;
    }}
  </style>
</head>
<body>
  <section class="page cover">
    <div class="cover-main">毕业论文（设计）</div>
    <div class="cover-title-label">论文（设计）题目：</div>
    <div class="cover-title">{esc(COVER['title_cn'])}</div>
    <div class="cover-title-en">{esc(COVER['title_en'])}</div>
    <table class="cover-table">
      <tr><td class="label">姓　　名</td><td class="value">{esc(COVER['name'])}</td></tr>
      <tr><td class="label">学　　号</td><td class="value">{esc(COVER['student_id'])}</td></tr>
      <tr><td class="label">学　　院</td><td class="value">{esc(COVER['college'])}</td></tr>
      <tr><td class="label">专　　业</td><td class="value">{esc(COVER['major'])}</td></tr>
      <tr><td class="label">年　　级</td><td class="value">{esc(COVER['grade'])}</td></tr>
      <tr><td class="label">指导教师</td><td class="value">{esc(COVER['advisor'])}</td></tr>
    </table>
    <div class="cover-date">{esc(COVER['date'])}</div>
  </section>

  <section class="page">
    <div class="grade-title">山东大学本科毕业论文（设计）成绩评定表</div>
    <table class="grade-table">
      <tr><td class="short">论文题目：{esc(COVER['title_cn'])}</td></tr>
      <tr><td>指导教师评语：</td></tr>
      <tr><td class="short">指导教师签名：　　　　　　　年　　月　　日</td></tr>
      <tr><td>评阅教师评语：</td></tr>
      <tr><td class="short">评阅教师签名：　　　　　　　年　　月　　日</td></tr>
      <tr><td>答辩小组评语：</td></tr>
      <tr><td class="short">最终成绩：　　　　　　　答辩小组组长签名：　　　　　　　年　　月　　日</td></tr>
    </table>
  </section>

  <section class="page abstract">
    <div class="abstract-title">摘　　要</div>
    {''.join(f'<p>{esc(p.strip())}</p>' for p in abstract_cn.splitlines() if p.strip())}
    <p class="keywords"><b>关键词：</b>{esc(keywords_cn)}</p>
  </section>

  <section class="page abstract">
    <div class="abstract-title">ABSTRACT</div>
    {''.join(f'<p>{esc(p.strip())}</p>' for p in EN_ABSTRACT.splitlines() if p.strip())}
    <p class="keywords"><b>Key Words:</b> {esc(EN_KEYWORDS)}</p>
  </section>

  <section class="page">
    <div class="toc-title">目　　录</div>
    {toc_html}
  </section>

  <main>
    {body_html}
  </main>

  <section>
    <div class="ref-title">参考文献</div>
    <ol class="references">
      {refs_html}
    </ol>
  </section>

  <section class="ack">
    <div class="ack-title">致　　谢</div>
    <p>在本论文完成过程中，感谢指导教师在选题论证、实验设计和论文修改方面给予的指导。感谢课题相关开源数据集、开源软件和研究社区提供的数据、工具与文献基础，使本文能够在 DrugComb 数据、图神经网络和 KAN 模型的基础上完成系统实验。感谢同学和朋友在实验调试、结果讨论和论文排版中提供的帮助。由于本人能力和时间有限，本文仍存在不足之处，恳请各位老师批评指正。</p>
  </section>
</body>
</html>
"""


def set_run_font(run, size=12, bold=False, name=EAST_ASIA_FONT, ascii_name: str | None = None) -> None:
    ascii_font = ascii_name or (LATIN_FONT if name in {EAST_ASIA_FONT, HEADING_FONT} else name)
    run.font.name = ascii_font
    if qn is not None and OxmlElement is not None:
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.rFonts
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.append(r_fonts)
        r_fonts.set(qn("w:eastAsia"), name)
        r_fonts.set(qn("w:ascii"), ascii_font)
        r_fonts.set(qn("w:hAnsi"), ascii_font)
        r_fonts.set(qn("w:cs"), ascii_font)
    run.font.size = Pt(size)
    run.font.bold = bold
    if RGBColor is not None:
        run.font.color.rgb = RGBColor(0, 0, 0)


def set_run_superscript(run) -> None:
    run.font.superscript = True
    if OxmlElement is not None and qn is not None:
        r_pr = run._element.get_or_add_rPr()
        vert_align = r_pr.find(qn("w:vertAlign"))
        if vert_align is None:
            vert_align = OxmlElement("w:vertAlign")
            r_pr.append(vert_align)
        vert_align.set(qn("w:val"), "superscript")


def set_run_subscript(run) -> None:
    run.font.subscript = True
    if OxmlElement is not None and qn is not None:
        r_pr = run._element.get_or_add_rPr()
        vert_align = r_pr.find(qn("w:vertAlign"))
        if vert_align is None:
            vert_align = OxmlElement("w:vertAlign")
            r_pr.append(vert_align)
        vert_align.set(qn("w:val"), "subscript")


def add_formatted_text_run(paragraph, text: str, size=12, bold=False, name=EAST_ASIA_FONT, superscript=False) -> None:
    if not text:
        return
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, name=name)
    if superscript:
        set_run_superscript(run)


def add_seq_field(paragraph, seq_name: str, result: str, size=10.5) -> None:
    if OxmlElement is None or qn is None:
        add_formatted_text_run(paragraph, result, size=size, name=LATIN_FONT)
        return

    run_begin = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run_begin._r.append(fld_begin)
    set_run_font(run_begin, size=size, name=LATIN_FONT)

    run_instr = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" SEQ {seq_name} \\* ARABIC "
    run_instr._r.append(instr)
    set_run_font(run_instr, size=size, name=LATIN_FONT)

    run_sep = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run_sep._r.append(fld_sep)
    set_run_font(run_sep, size=size, name=LATIN_FONT)

    run_result = paragraph.add_run(result)
    set_run_font(run_result, size=size, name=LATIN_FONT)

    run_end = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run_end._r.append(fld_end)
    set_run_font(run_end, size=size, name=LATIN_FONT)


def add_runs_with_citations(paragraph, text: str, size=12, bold=False, name=EAST_ASIA_FONT) -> None:
    pos = 0
    for match in CITATION_RE.finditer(text):
        if match.start() > pos:
            add_formatted_text_run(
                paragraph,
                text[pos:match.start()],
                size=size,
                bold=bold,
                name=name,
            )
        add_internal_hyperlink(
            paragraph,
            match.group(0),
            anchor=citation_anchor(match.group(0)),
            size=max(size - 1, 8),
            bold=bold,
            name=LATIN_FONT,
            superscript=True,
        )
        pos = match.end()
    if pos < len(text):
        add_formatted_text_run(paragraph, text[pos:], size=size, bold=bold, name=name)


def set_section_layout(section) -> None:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)


def set_page_number_format(section, fmt: str, start: int = 1) -> None:
    if OxmlElement is None or qn is None:
        return
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(start))
    pg_num_type.set(qn("w:fmt"), fmt)


def add_page_number_field(paragraph) -> None:
    if OxmlElement is None or qn is None:
        return
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, size=9, name=LATIN_FONT)


_BOOKMARK_ID = 0


def next_bookmark_id() -> int:
    global _BOOKMARK_ID
    _BOOKMARK_ID += 1
    return _BOOKMARK_ID


def add_bookmarked_run(paragraph, text: str, bookmark: str | None = None, size=12, bold=False, name=EAST_ASIA_FONT):
    if bookmark and OxmlElement is not None and qn is not None:
        bookmark_id = next_bookmark_id()
        start = OxmlElement("w:bookmarkStart")
        start.set(qn("w:id"), str(bookmark_id))
        start.set(qn("w:name"), bookmark)
        paragraph._p.append(start)
        run = paragraph.add_run(text)
        set_run_font(run, size=size, bold=bold, name=name)
        end = OxmlElement("w:bookmarkEnd")
        end.set(qn("w:id"), str(bookmark_id))
        paragraph._p.append(end)
        return run
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, name=name)
    return run


def add_internal_hyperlink(
    paragraph,
    text: str,
    anchor: str,
    size=12,
    bold=False,
    name=EAST_ASIA_FONT,
    superscript=False,
) -> None:
    if OxmlElement is None or qn is None:
        run = paragraph.add_run(text)
        set_run_font(run, size=size, bold=bold, name=name)
        if superscript:
            set_run_superscript(run)
        return

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    r_fonts = OxmlElement("w:rFonts")
    ascii_font = LATIN_FONT if name in {EAST_ASIA_FONT, HEADING_FONT} else name
    r_fonts.set(qn("w:eastAsia"), name)
    for key in ("w:ascii", "w:hAnsi", "w:cs"):
        r_fonts.set(qn(key), ascii_font)
    r_pr.append(r_fonts)

    if bold:
        r_pr.append(OxmlElement("w:b"))
        r_pr.append(OxmlElement("w:bCs"))
    if superscript:
        vert_align = OxmlElement("w:vertAlign")
        vert_align.set(qn("w:val"), "superscript")
        r_pr.append(vert_align)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "none")
    r_pr.append(underline)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    r_pr.append(sz)
    sz_cs = OxmlElement("w:szCs")
    sz_cs.set(qn("w:val"), str(int(size * 2)))
    r_pr.append(sz_cs)

    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(r_pr)
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def citation_anchor(citation_text: str) -> str:
    match = re.match(r"\[(\d+)", citation_text)
    return f"bm_ref_{match.group(1)}" if match else "bm_refs"


def collect_toc_entries(body_md: str) -> list[tuple[str, int, str]]:
    entries: list[tuple[str, int, str]] = []
    idx = 1
    for line in body_md.splitlines():
        if line.startswith("## "):
            title = line[3:].strip()
            entries.append((title, 1, f"bm_sec_{idx}"))
            idx += 1
        elif line.startswith("### "):
            title = line[4:].strip()
            entries.append((title, 2, f"bm_sec_{idx}"))
            idx += 1
        elif line.startswith("#### "):
            title = line[5:].strip()
            entries.append((title, 3, f"bm_sec_{idx}"))
            idx += 1
    entries.append(("参考文献", 1, "bm_refs"))
    entries.append(("致谢", 1, "bm_ack"))
    return entries


INLINE_MATH_RE = re.compile(r"(?:\\\\|\\)\((.*?)(?:\\\\|\\)\)")
LITERAL_LBRACE = "§LBRACE§"
LITERAL_RBRACE = "§RBRACE§"


def _subscript(text: str) -> str:
    chars = {
        "0": "₀", "1": "₁", "2": "₂", "3": "₃", "4": "₄",
        "5": "₅", "6": "₆", "7": "₇", "8": "₈", "9": "₉",
        "+": "₊", "-": "₋", "=": "₌", "(": "₍", ")": "₎",
        "a": "ₐ", "e": "ₑ", "h": "ₕ", "i": "ᵢ", "j": "ⱼ",
        "k": "ₖ", "l": "ₗ", "m": "ₘ", "n": "ₙ", "o": "ₒ",
        "p": "ₚ", "r": "ᵣ", "s": "ₛ", "t": "ₜ", "u": "ᵤ",
        "v": "ᵥ", "x": "ₓ", ",": ",",
    }
    if all(ch in chars for ch in text):
        return "".join(chars[ch] for ch in text)
    return "_" + text


def _superscript(text: str) -> str:
    chars = {
        "0": "⁰", "1": "¹", "2": "²", "3": "³", "4": "⁴",
        "5": "⁵", "6": "⁶", "7": "⁷", "8": "⁸", "9": "⁹",
        "+": "⁺", "-": "⁻", "=": "⁼", "(": "⁽", ")": "⁾",
        "a": "ᵃ", "b": "ᵇ", "c": "ᶜ", "d": "ᵈ", "e": "ᵉ",
        "f": "ᶠ", "g": "ᵍ", "h": "ʰ", "i": "ⁱ", "j": "ʲ",
        "k": "ᵏ", "l": "ˡ", "m": "ᵐ", "n": "ⁿ", "o": "ᵒ",
        "p": "ᵖ", "r": "ʳ", "s": "ˢ", "t": "ᵗ", "u": "ᵘ",
        "v": "ᵛ", "w": "ʷ", "x": "ˣ", "y": "ʸ", "z": "ᶻ",
        "A": "ᴬ", "B": "ᴮ", "D": "ᴰ", "E": "ᴱ", "G": "ᴳ",
        "H": "ᴴ", "I": "ᴵ", "J": "ᴶ", "K": "ᴷ", "L": "ᴸ",
        "M": "ᴹ", "N": "ᴺ", "O": "ᴼ", "P": "ᴾ", "R": "ᴿ",
        "T": "ᵀ", "U": "ᵁ", "V": "ⱽ", "W": "ᵂ",
    }
    if all(ch in chars for ch in text):
        return "".join(chars[ch] for ch in text)
    return "^" + text


def normalize_latex_math(expr: str) -> str:
    s = expr.replace("\\\\", "\\").strip()
    s = s.replace("\\{", LITERAL_LBRACE).replace("\\}", LITERAL_RBRACE)
    s = re.sub(r"\\operatorname\{([^{}]+)\}", r"\1", s)
    s = re.sub(r"\\mathbb\{R\}", "ℝ", s)
    s = re.sub(r"\\mathcal\{N\}", "𝒩", s)
    s = re.sub(r"\\frac\{([^{}]+)\}\{([^{}]+)\}", r"\1/\2", s)
    s = re.sub(r"\\hat\{([^{}]+)\}", lambda m: m.group(1) + "\u0302", s)
    s = re.sub(r"\\bar\{([^{}]+)\}", lambda m: m.group(1) + "\u0304", s)
    replacements = {
        "\\alpha": "α",
        "\\phi": "φ",
        "\\pi": "π",
        "\\in": "∈",
        "\\sum": "Σ",
        "\\cos": "cos",
        "\\sin": "sin",
        "\\left": "",
        "\\right": "",
    }
    for old, new in replacements.items():
        s = s.replace(old, new)
    s = re.sub(r"\\([A-Za-z]+)", r"\1", s)
    s = re.sub(r"\s+", " ", s)
    return s.strip()


def restore_math_text(text: str) -> str:
    text = text.replace("{", "").replace("}", "")
    return text.replace(LITERAL_LBRACE, "{").replace(LITERAL_RBRACE, "}")


def read_math_script_group(text: str, start: int) -> tuple[str, int]:
    if start >= len(text):
        return "", start
    if text[start] != "{":
        return text[start], start + 1

    depth = 1
    pos = start + 1
    group: list[str] = []
    while pos < len(text):
        char = text[pos]
        if char == "{":
            depth += 1
        elif char == "}":
            depth -= 1
            if depth == 0:
                return "".join(group), pos + 1
        group.append(char)
        pos += 1
    return "".join(group), pos


def parse_math_tokens(expr: str) -> list[tuple[str, str]]:
    s = normalize_latex_math(expr)
    tokens: list[tuple[str, str]] = []
    buf: list[str] = []
    pos = 0

    def flush() -> None:
        if buf:
            text = restore_math_text("".join(buf))
            if text:
                tokens.append((text, "normal"))
            buf.clear()

    while pos < len(s):
        char = s[pos]
        if char in {"_", "^"}:
            flush()
            group, next_pos = read_math_script_group(s, pos + 1)
            text = restore_math_text(group)
            if text:
                tokens.append((text, "subscript" if char == "_" else "superscript"))
            pos = next_pos
            continue
        buf.append(char)
        pos += 1
    flush()
    return tokens


def clean_latex_math(expr: str) -> str:
    return "".join(text for text, _ in parse_math_tokens(expr)).strip()


def render_math_html(expr: str) -> str:
    parts: list[str] = []
    for text, kind in parse_math_tokens(expr):
        if kind == "subscript":
            parts.append(f"<sub>{esc(text)}</sub>")
        elif kind == "superscript":
            parts.append(f"<sup>{esc(text)}</sup>")
        else:
            parts.append(esc(text))
    return f"<span class='math'>{''.join(parts)}</span>"


def add_runs_with_inline_math(paragraph, text: str, size=12, bold=False, name=EAST_ASIA_FONT) -> None:
    pos = 0
    for match in INLINE_MATH_RE.finditer(text):
        if match.start() > pos:
            add_runs_with_citations(
                paragraph,
                text[pos:match.start()],
                size=size,
                bold=bold,
                name=name,
            )
        add_math_runs(paragraph, match.group(1), size=size, bold=bold)
        pos = match.end()
    if pos < len(text):
        add_runs_with_citations(paragraph, text[pos:], size=size, bold=bold, name=name)


def add_math_runs(paragraph, expr: str, size=12, bold=False) -> None:
    for token, kind in parse_math_tokens(expr):
        run = paragraph.add_run(token)
        set_run_font(run, size=size, bold=bold, name=MATH_FONT)
        if kind == "subscript":
            set_run_subscript(run)
        elif kind == "superscript":
            set_run_superscript(run)


def is_math_block_start(text: str) -> bool:
    return text in {"\\\\[", "\\["}


def is_math_block_end(text: str) -> bool:
    return text in {"\\\\]", "\\]"}


def add_equation_docx(doc, expr: str, number: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(3)
    if WD_TAB_ALIGNMENT is not None:
        p.paragraph_format.tab_stops.add_tab_stop(Cm(7.5), WD_TAB_ALIGNMENT.CENTER)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(15.0), WD_TAB_ALIGNMENT.RIGHT)
        run = p.add_run("\t")
        set_run_font(run, size=12, name=MATH_FONT)
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_math_runs(p, expr, size=12)
    if number:
        run = p.add_run("\t" + number)
        set_run_font(run, size=12, name=EAST_ASIA_FONT)


def add_center_page_number_footer(section, fmt: str = "decimal", start: int = 1) -> None:
    section.footer.is_linked_to_previous = False
    set_page_number_format(section, fmt=fmt, start=start)
    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number_field(footer_p)


def set_paragraph_body(paragraph, first_line=True) -> None:
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(3)
    if first_line:
        paragraph.paragraph_format.first_line_indent = Cm(0.74)


def add_text_paragraph(doc, text, first_line=True, align=None, size=12, bold=False, name=EAST_ASIA_FONT) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    set_paragraph_body(p, first_line=first_line)
    add_runs_with_inline_math(p, text, size=size, bold=bold, name=name)


def add_keyword_paragraph(
    doc,
    label: str,
    content: str,
    label_name=HEADING_FONT,
    content_name=EAST_ASIA_FONT,
) -> None:
    p = doc.add_paragraph()
    set_paragraph_body(p, first_line=False)
    add_formatted_text_run(p, label, size=12, bold=True, name=label_name)
    add_runs_with_inline_math(p, content, size=12, bold=False, name=content_name)


def add_reference_paragraph(doc, text: str, bookmark: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(0)
    add_bookmarked_run(p, text, bookmark=bookmark, size=10.5, name=EAST_ASIA_FONT)


def is_caption_line(text: str, prefix: str | None = None) -> bool:
    match = CAPTION_RE.match(text.strip())
    if not match:
        return False
    return prefix is None or match.group(1) == prefix


def next_nonempty_line(lines: list[str], start: int) -> str:
    for idx in range(start, len(lines)):
        candidate = lines[idx].strip()
        if candidate:
            return candidate
    return ""


def is_markdown_block_boundary(text: str, lines: list[str] | None = None, idx: int = 0) -> bool:
    if not text:
        return True
    if text.startswith(("## ", "### ", "#### ", "|", "![", "- ")):
        return True
    if is_math_block_start(text) or is_caption_line(text):
        return True
    if re.match(r"\d+\.\s+", text):
        return True
    return False


def ensure_note_prefix(text: str) -> str:
    stripped = text.strip()
    if not stripped:
        return stripped
    if stripped.startswith("注："):
        return stripped
    if stripped.startswith("注:"):
        return "注：" + stripped[2:].lstrip()
    return "注：" + stripped


def add_caption_docx(doc, text: str, kind: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(3 if kind == "表" else 2)
    p.paragraph_format.space_after = Pt(4 if kind == "图" else 3)
    p.paragraph_format.keep_together = True
    if kind == "表":
        p.paragraph_format.keep_with_next = True
    match = re.match(r"^(图|表)\s*(\d+)\s*(.*)$", text.strip())
    if match:
        caption_kind, number, rest = match.groups()
        add_formatted_text_run(p, f"{caption_kind} ", size=10.5, bold=True, name=EAST_ASIA_FONT)
        add_seq_field(
            p,
            "Figure" if caption_kind == "图" else "Table",
            result=number,
            size=10.5,
        )
        for run in p.runs:
            run.font.bold = True
        add_runs_with_inline_math(p, f" {rest}".rstrip(), size=10.5, bold=True, name=EAST_ASIA_FONT)
    else:
        run = p.add_run(text.strip())
        set_run_font(run, size=10.5, bold=True, name=EAST_ASIA_FONT)


def add_figure_note_docx(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.left_indent = Cm(0.74)
    p.paragraph_format.right_indent = Cm(0.74)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_together = True
    add_runs_with_inline_math(p, ensure_note_prefix(text), size=10.5, name=EAST_ASIA_FONT)


def add_center_title(doc, text, size=18, bold=True, bookmark: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    add_bookmarked_run(p, text, bookmark=bookmark, size=size, bold=bold, name=HEADING_FONT)


def add_center_heading_title(doc, text, size=18, bold=True, bookmark: str | None = None) -> None:
    p = doc.add_paragraph(style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    add_bookmarked_run(p, text, bookmark=bookmark, size=size, bold=bold, name=HEADING_FONT)


def add_heading_docx(doc, text, level=1, bookmark: str | None = None) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(6)
    size_map = {1: 16, 2: 14, 3: 12}
    add_bookmarked_run(
        p,
        text,
        bookmark=bookmark,
        size=size_map.get(level, 12),
        bold=True,
        name=HEADING_FONT,
    )


def add_cover_docx(doc) -> None:
    for section in doc.sections:
        set_section_layout(section)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Cm(1.5)
    p.paragraph_format.space_after = Cm(2.0)
    run = p.add_run("毕 业 论 文（设 计）")
    set_run_font(run, size=26, bold=True, name=HEADING_FONT)

    add_text_paragraph(doc, "论文（设计）题目：", first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    add_text_paragraph(doc, COVER["title_cn"], first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=18, bold=True, name=HEADING_FONT)
    add_text_paragraph(doc, COVER["title_en"], first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=16, bold=True, name=LATIN_FONT)

    table = doc.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER if WD_TABLE_ALIGNMENT is not None else 1
    table.autofit = False
    fields = [
        ("姓    名", COVER["name"]),
        ("学    号", COVER["student_id"]),
        ("学    院", COVER["college"]),
        ("专    业", COVER["major"]),
        ("年    级", COVER["grade"]),
        ("指导教师", COVER["advisor"]),
    ]
    for row, (label, value) in zip(table.rows, fields):
        row.height = Cm(0.85)
        row.cells[0].text = label
        row.cells[1].text = value
        row.cells[0].width = Cm(3.1)
        row.cells[1].width = Cm(7.6)
        set_cell_border(row.cells[0], top="nil", left="nil", bottom="nil", right="nil")
        set_cell_border(row.cells[1], top="nil", left="nil", bottom="single", right="nil")
        for cell in row.cells:
            if WD_CELL_VERTICAL_ALIGNMENT is not None:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_run_font(run, size=14, name=EAST_ASIA_FONT)

    add_text_paragraph(doc, COVER["date"], first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    doc.add_page_break()


def add_grade_form_docx(doc) -> None:
    add_center_title(doc, "山东大学本科毕业论文（设计）成绩评定表")
    table = doc.add_table(rows=7, cols=1)
    table.style = "Table Grid"
    rows = [
        f"论文题目：{COVER['title_cn']}",
        "指导教师评语：",
        "指导教师签名：　　　　　　　年　　月　　日",
        "评阅教师评语：",
        "评阅教师签名：　　　　　　　年　　月　　日",
        "答辩小组评语：",
        "最终成绩：　　　　　　　答辩小组组长签名：　　　　　　　年　　月　　日",
    ]
    for row, text in zip(table.rows, rows):
        cell = row.cells[0]
        cell.text = text
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.line_spacing = 1.5
            for run in paragraph.runs:
                set_run_font(run, size=12)


def add_abstract_docx(doc, abstract_cn, keywords_cn) -> None:
    add_center_title(doc, "摘    要", bookmark="bm_abs_cn")
    for para in abstract_cn.splitlines():
        para = para.strip()
        if para:
            add_text_paragraph(doc, para)
    add_keyword_paragraph(doc, "关键词：", keywords_cn, label_name=HEADING_FONT, content_name=EAST_ASIA_FONT)
    doc.add_page_break()

    add_center_title(doc, "ABSTRACT", bookmark="bm_abs_en")
    for para in EN_ABSTRACT.splitlines():
        para = para.strip()
        if para:
            add_text_paragraph(doc, para)
    add_keyword_paragraph(doc, "Key Words: ", EN_KEYWORDS, label_name=LATIN_FONT, content_name=LATIN_FONT)
    doc.add_page_break()


def add_toc_docx(doc, toc_entries, toc_pages=None) -> None:
    add_center_title(doc, "目    录")
    for title, level, anchor in toc_entries:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Cm({1: 0, 2: 0.74, 3: 1.48}.get(level, 0))
        if WD_TAB_ALIGNMENT is not None and WD_TAB_LEADER is not None:
            p.paragraph_format.tab_stops.add_tab_stop(
                Cm(14.8),
                WD_TAB_ALIGNMENT.RIGHT,
                WD_TAB_LEADER.DOTS,
            )
        add_internal_hyperlink(p, title, anchor, size=12, bold=(level == 1))
        if toc_pages and anchor in toc_pages:
            run = p.add_run("\t")
            set_run_font(run, size=12)
            add_internal_hyperlink(p, str(toc_pages[anchor]), anchor, size=12, bold=False)
    doc.add_page_break()


def set_cell_border(cell, **kwargs) -> None:
    if OxmlElement is None or qn is None:
        return
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        if edge not in kwargs:
            continue
        tag = "w:" + edge
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        val = kwargs[edge]
        element.set(qn("w:val"), val)
        if val != "nil":
            element.set(qn("w:sz"), "8")
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), "000000")


def apply_three_line_table(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top="nil", left="nil", bottom="nil", right="nil")
    if not table.rows:
        return
    for cell in table.rows[0].cells:
        set_cell_border(cell, top="single", bottom="single", left="nil", right="nil")
    for cell in table.rows[-1].cells:
        set_cell_border(cell, bottom="single", left="nil", right="nil")


def set_table_row_cant_split(row) -> None:
    if OxmlElement is None or qn is None:
        return
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_table_row_header(row) -> None:
    if OxmlElement is None or qn is None:
        return
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)


def set_cell_width(cell, width_cm: float) -> None:
    cell.width = Cm(width_cm)
    if OxmlElement is None or qn is None:
        return
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def table_col_widths(col_count: int) -> list[float]:
    # Keep tables inside the 15 cm text area while giving model-name columns room.
    if col_count == 3:
        return [3.2, 6.0, 5.8]
    if col_count == 5:
        return [3.2, 2.8, 3.0, 2.8, 3.2]
    if col_count == 6:
        return [3.0, 3.2, 2.2, 2.2, 2.2, 2.2]
    return [15.0 / col_count] * col_count


def add_table_docx(doc, lines) -> None:
    rows = []
    for line in lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", c or "") for c in cells):
            continue
        rows.append(cells)
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER if WD_TABLE_ALIGNMENT is not None else 1
    table.autofit = False
    col_widths = table_col_widths(max(len(r) for r in rows))
    apply_three_line_table(table)
    for i, row in enumerate(rows):
        table.rows[i].height = Cm(0.55)
        set_table_row_cant_split(table.rows[i])
        if i == 0:
            set_table_row_header(table.rows[i])
        for j, cell_text in enumerate(row):
            cell = table.cell(i, j)
            if WD_CELL_VERTICAL_ALIGNMENT is not None:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_width(cell, col_widths[j] if j < len(col_widths) else 15.0 / len(col_widths))
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.keep_together = True
                paragraph.paragraph_format.keep_with_next = i < len(rows) - 1
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    set_run_font(run, size=10.5, bold=(i == 0), name=EAST_ASIA_FONT)


def add_body_markdown_docx(doc, md: str, heading_entries) -> None:
    lines = md.splitlines()
    i = 0
    paragraph = []
    heading_idx = 0
    current_chapter = 0
    equation_counts: dict[int, int] = {}
    pending_table_caption: str | None = None
    pending_figure = False
    expect_figure_note = False

    def flush_paragraph() -> None:
        if paragraph:
            add_text_paragraph(doc, " ".join(x.strip() for x in paragraph).strip())
            paragraph.clear()

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            flush_paragraph()
            i += 1
            continue

        if is_caption_line(stripped, "表") and next_nonempty_line(lines, i + 1).startswith("|"):
            flush_paragraph()
            pending_table_caption = stripped
            i += 1
            continue

        if is_caption_line(stripped, "图") and pending_figure:
            flush_paragraph()
            add_caption_docx(doc, stripped, kind="图")
            pending_figure = False
            expect_figure_note = True
            i += 1
            continue

        if expect_figure_note:
            if is_markdown_block_boundary(stripped, lines, i):
                expect_figure_note = False
            else:
                flush_paragraph()
                note_lines = []
                while i < len(lines):
                    candidate = lines[i].strip()
                    if not candidate or is_markdown_block_boundary(candidate, lines, i):
                        break
                    note_lines.append(lines[i])
                    i += 1
                add_figure_note_docx(doc, " ".join(x.strip() for x in note_lines).strip())
                expect_figure_note = False
                continue

        if stripped.startswith("|"):
            flush_paragraph()
            if pending_table_caption:
                add_caption_docx(doc, pending_table_caption, kind="表")
                pending_table_caption = None
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_table_docx(doc, table_lines)
            continue

        image_match = re.match(r"!\[(.*?)\]\((.*?)\)", stripped)
        if image_match:
            flush_paragraph()
            _, rel = image_match.groups()
            image_path = THESIS_DIR / rel
            if image_path.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(1)
                run = p.add_run()
                run.add_picture(str(image_path), width=Inches(5.9))
                pending_figure = True
            i += 1
            continue

        if is_math_block_start(stripped):
            flush_paragraph()
            equation_lines = []
            i += 1
            while i < len(lines):
                if is_math_block_end(lines[i].strip()):
                    i += 1
                    break
                equation_lines.append(lines[i].strip())
                i += 1
            chapter = current_chapter or 0
            equation_counts[chapter] = equation_counts.get(chapter, 0) + 1
            if current_chapter:
                number = f"({current_chapter}-{equation_counts[chapter]})"
            else:
                number = f"({equation_counts[chapter]})"
            add_equation_docx(doc, " ".join(equation_lines), number=number)
            continue

        if stripped.startswith("#### "):
            flush_paragraph()
            bookmark = heading_entries[heading_idx][2] if heading_idx < len(heading_entries) else None
            heading_idx += 1
            add_heading_docx(doc, stripped[5:].strip(), level=3, bookmark=bookmark)
            i += 1
            continue

        if stripped.startswith("### "):
            flush_paragraph()
            bookmark = heading_entries[heading_idx][2] if heading_idx < len(heading_entries) else None
            heading_idx += 1
            add_heading_docx(doc, stripped[4:].strip(), level=2, bookmark=bookmark)
            i += 1
            continue

        if stripped.startswith("## "):
            flush_paragraph()
            bookmark = heading_entries[heading_idx][2] if heading_idx < len(heading_entries) else None
            heading_idx += 1
            match = re.match(r"##\s+(\d+)", stripped)
            if match:
                current_chapter = int(match.group(1))
            add_heading_docx(doc, stripped[3:].strip(), level=1, bookmark=bookmark)
            i += 1
            continue

        if stripped.startswith("- "):
            flush_paragraph()
            while i < len(lines) and lines[i].strip().startswith("- "):
                add_text_paragraph(doc, "· " + lines[i].strip()[2:], first_line=False)
                i += 1
            continue

        if re.match(r"\d+\.\s+", stripped):
            flush_paragraph()
            while i < len(lines) and re.match(r"\d+\.\s+", lines[i].strip()):
                add_text_paragraph(doc, lines[i].strip(), first_line=False)
                i += 1
            continue

        paragraph.append(line)
        i += 1

    flush_paragraph()


def build_docx(toc_pages=None) -> bool:
    global _BOOKMARK_ID
    _BOOKMARK_ID = 0
    if Document is None:
        return False
    draft_text = DRAFT.read_text(encoding="utf-8")
    _, abstract_cn, keywords_cn, body_md = split_draft(draft_text)
    references = extract_references(draft_text)
    toc_entries = collect_toc_entries(body_md)
    heading_entries = [entry for entry in toc_entries if entry[2].startswith("bm_sec_")]
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "SimSun"
    styles["Normal"].font.size = Pt(12)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

    add_cover_docx(doc)
    add_grade_form_docx(doc)

    abstract_section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_section_layout(abstract_section)
    abstract_section.footer.is_linked_to_previous = False
    add_abstract_docx(doc, abstract_cn, keywords_cn)

    toc_section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_section_layout(toc_section)
    add_center_page_number_footer(toc_section, fmt="upperRoman", start=1)
    add_toc_docx(doc, toc_entries, toc_pages=toc_pages)

    section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_section_layout(section)
    section.header.is_linked_to_previous = False
    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_p.add_run("山东大学本科毕业论文（设计）")
    set_run_font(header_run, size=9)
    add_center_page_number_footer(section, fmt="decimal", start=1)

    add_body_markdown_docx(doc, body_md, heading_entries)
    add_center_heading_title(doc, "参考文献", bookmark="bm_refs")
    for i, ref in enumerate(references, 1):
        add_reference_paragraph(doc, f"[{i}] {ref}", bookmark=f"bm_ref_{i}")

    doc.add_page_break()
    add_center_heading_title(doc, "致    谢", bookmark="bm_ack")
    add_text_paragraph(
        doc,
        "在本论文完成过程中，感谢指导教师在选题论证、实验设计和论文修改方面给予的指导。感谢课题相关开源数据集、开源软件和研究社区提供的数据、工具与文献基础，使本文能够在 DrugComb 数据、图神经网络和 KAN 模型的基础上完成系统实验。感谢同学和朋友在实验调试、结果讨论和论文排版中提供的帮助。由于本人能力和时间有限，本文仍存在不足之处，恳请各位老师批评指正。",
    )

    doc.save(OUT_DOCX)
    return True


def collect_docx_bookmark_pages(docx_path: Path, toc_entries) -> dict[str, int]:
    try:
        import time
        import uno
        from com.sun.star.beans import PropertyValue
    except Exception:
        return {}

    port = 2019
    proc = subprocess.Popen(
        [
            "libreoffice",
            "--headless",
            f"--accept=socket,host=localhost,port={port};urp;StarOffice.ServiceManager",
            "--norestore",
            "--nofirststartwizard",
        ],
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
    )
    try:
        local_ctx = uno.getComponentContext()
        resolver = local_ctx.ServiceManager.createInstanceWithContext(
            "com.sun.star.bridge.UnoUrlResolver",
            local_ctx,
        )
        ctx = None
        for _ in range(20):
            try:
                ctx = resolver.resolve(
                    f"uno:socket,host=localhost,port={port};urp;StarOffice.ComponentContext"
                )
                break
            except Exception:
                time.sleep(0.25)
        if ctx is None:
            return {}

        smgr = ctx.ServiceManager
        desktop = smgr.createInstanceWithContext("com.sun.star.frame.Desktop", ctx)
        props = []
        for name, value in [("Hidden", True), ("ReadOnly", True)]:
            prop = PropertyValue()
            prop.Name = name
            prop.Value = value
            props.append(prop)
        doc = desktop.loadComponentFromURL(docx_path.resolve().as_uri(), "_blank", 0, tuple(props))
        controller = doc.getCurrentController()
        view_cursor = controller.getViewCursor()
        bookmarks = doc.getBookmarks()
        pages: dict[str, int] = {}
        for _, _, anchor in toc_entries:
            try:
                bookmark = bookmarks.getByName(anchor)
                view_cursor.gotoRange(bookmark.getAnchor(), False)
                pages[anchor] = int(view_cursor.Page)
            except Exception:
                continue
        doc.close(True)
        return pages
    finally:
        proc.terminate()
        try:
            proc.wait(timeout=5)
        except Exception:
            proc.kill()


def to_roman(value: int) -> str:
    pairs = [
        (1000, "M"), (900, "CM"), (500, "D"), (400, "CD"),
        (100, "C"), (90, "XC"), (50, "L"), (40, "XL"),
        (10, "X"), (9, "IX"), (5, "V"), (4, "IV"), (1, "I"),
    ]
    out = []
    remaining = value
    for number, numeral in pairs:
        while remaining >= number:
            out.append(numeral)
            remaining -= number
    return "".join(out)


def build_docx_with_toc_pages() -> bool:
    if not build_docx():
        return False
    draft_text = DRAFT.read_text(encoding="utf-8")
    _, _, _, body_md = split_draft(draft_text)
    toc_entries = collect_toc_entries(body_md)
    pages = collect_docx_bookmark_pages(OUT_DOCX, toc_entries)
    if pages:
        front_start = pages.get("bm_abs_cn", 1)
        body_start = pages.get("bm_sec_1", 1)
        display_pages: dict[str, str | int] = {}
        for _, _, anchor in toc_entries:
            if anchor not in pages:
                continue
            if anchor in {"bm_abs_cn", "bm_abs_en"}:
                display_pages[anchor] = to_roman(max(1, pages[anchor] - front_start + 1))
            else:
                display_pages[anchor] = max(1, pages[anchor] - body_start + 1)
        return build_docx(toc_pages=display_pages)
    return True


def convert_with_libreoffice() -> None:
    if not build_docx_with_toc_pages():
        subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "odt", "--outdir", str(THESIS_DIR), str(OUT_HTML)],
            check=False,
        )
        if OUT_ODT.exists():
            subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "docx", "--outdir", str(THESIS_DIR), str(OUT_ODT)],
                check=False,
            )
    if OUT_DOCX.exists():
        subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", str(THESIS_DIR), str(OUT_DOCX)],
            check=False,
        )


def main() -> None:
    OUT_HTML.write_text(build_html(), encoding="utf-8")
    convert_with_libreoffice()
    print(f"Wrote {OUT_HTML}")
    if OUT_DOCX.exists():
        print(f"Wrote {OUT_DOCX}")
    if OUT_PDF.exists():
        print(f"Wrote {OUT_PDF}")


if __name__ == "__main__":
    main()
